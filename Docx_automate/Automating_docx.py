#! python3

import docx
from docx.shared import Inches

d=docx.Document()

def title(text):
   p=d.add_heading(text)
   p.style='Title'

def heading(text):
   d.add_heading(text,level=1)

def picture(location,inches):
   d.add_picture(location,width=Inches(inches))

def list_(list1,type_):
   for nlist in list1:
      if type_=='ordered':                                
         d.add_paragraph(nlist, style='List Number')      
      elif type_=='unordered':                                   
         d.add_paragraph(nlist, style='List Bullet')
    
def paragraphs(text,confirm_,style_,tnum):
   p=d.add_paragraph(text)
   if confirm_=='yes':
      if style_=='italic':
         p.runs[tnum].italic=True
      elif style_=='bold':
         p.runs[tnum].bold=True
      elif style_=='underline':
         p.runs[tnum].underline=True   

   
def save(filename):
   d.save(filename+'.docx')

while True:   
   fname=input('Enter the operation name: ')
##   text=input('Enter the text you want to insert')
##   fname+'('+str(text)+')'         #-- logic error--- I am just concatenating string not calling the function
   if fname =='title':
      text=input('Enter the title you want to insert: ')
      title(text)
   elif fname=='heading':
      text=input('Enter the heading you want to insert: ')
      heading(text)
   elif fname=='picture':
      text=input('Enter the picture you want to insert: ')
      inch=int(input('Specify the inch of the photo:'))
      picture(text,inch)
   elif fname=='list':
      type_=input('Enter the type of list: ')
      number=int(input('how many points you need have on list: '))
      list1_=[]
      for rep in range(number):
         text=input('Enter the text you want to insert: ')
         list1_.append(text)
      list_(list1_,type_)
   elif fname=='paragraphs':
      text=input('Enter the paragraph you want to insert: ')
      confirm_=input('Do you want to apply style to your paragraph:')
      tnum=int(input('Which text you want to style it:'))
      style_=input('Enter the style:')
      paragraphs(text,confirm_,style_,tnum)
   elif fname=='save':
      filename=input('Enter the filename:')
      save(filename)   
   elif fname=='close':
      break




   
  
